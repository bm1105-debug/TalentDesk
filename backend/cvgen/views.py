# cvgen/views.py
# What this file does: two read-only endpoints that stream a generated
# PDF or DOCX file for a given candidate. No data is persisted.

from io import BytesIO

from django.http import HttpResponse
from django.template.loader import render_to_string
from django.shortcuts import get_object_or_404

from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated

from xhtml2pdf import pisa                       # HTML → PDF renderer
from docx import Document                        # python-docx DOCX builder
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from candidates.models import Candidate
from users.permissions import IsRecruiterOrAbove


# ── Shared helper ──────────────────────────────────────────────────────────────

def _get_candidate(pk):
    """Fetch candidate with skills pre-fetched — both views need the same data."""
    return get_object_or_404(
        Candidate.objects.prefetch_related("skills"), pk=pk
    )

# ── PDF View ───────────────────────────────────────────────────────────────────

class CandidatePDFView(APIView):
    """
    GET /api/cvgen/candidates/{id}/pdf/
    Returns a downloadable PDF built from the cv.html template.
    """
    permission_classes = [IsRecruiterOrAbove]

    def get(self, request, pk):
        candidate = _get_candidate(pk)

        # Render the Django HTML template into a string
        html_string = render_to_string("cvgen/cv.html", {
            "candidate": candidate,
            "skills":    candidate.skills.all(),
        })

        # Convert the HTML string to PDF bytes in memory
        buffer = BytesIO()
        result = pisa.CreatePDF(html_string, dest=buffer)

        if result.err:
            # pisa sets result.err > 0 on render failure
            return HttpResponse("PDF generation failed.", status=500)

        filename = f"{candidate.first_name}_{candidate.last_name}_CV.pdf"

        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        # 'attachment' prompts a Save-As dialog in the browser
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


# ── DOCX View ──────────────────────────────────────────────────────────────────

class CandidateDOCXView(APIView):
    """
    GET /api/cvgen/candidates/{id}/docx/
    Returns a downloadable Word document built with python-docx.
    """
    permission_classes = [IsRecruiterOrAbove]

    def get(self, request, pk):
        candidate = _get_candidate(pk)
        skills    = list(candidate.skills.values_list("name", flat=True))

        doc = Document()

        # ── Page margins (narrow for a CV feel) ──
        for section in doc.sections:
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin   = Inches(0.9)
            section.right_margin  = Inches(0.9)

        # ── Name heading ──
        name_para = doc.add_paragraph()
        name_run  = name_para.add_run(
            f"{candidate.first_name} {candidate.last_name}"
        )
        name_run.bold      = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)  # brand blue

        # ── Current title + company ──
        if candidate.current_title or candidate.current_company:
            parts = filter(None, [candidate.current_title, candidate.current_company])
            sub = doc.add_paragraph(" · ".join(parts))
            sub.runs[0].font.size = Pt(11)
            sub.runs[0].italic    = True

        # ── Contact line ──
        contact_parts = [candidate.email, candidate.phone]
        if candidate.location:    contact_parts.append(candidate.location)
        if candidate.linkedin_url: contact_parts.append(candidate.linkedin_url)
        contact_para = doc.add_paragraph("  |  ".join(contact_parts))
        contact_para.runs[0].font.size = Pt(9)

        doc.add_paragraph()   # blank spacer

        # ── Skills section ──
        if skills:
            self._section_heading(doc, "Skills")
            doc.add_paragraph(", ".join(skills))
            doc.add_paragraph()

        # ── Summary / notes section ──
        if candidate.notes:
            self._section_heading(doc, "Summary")
            doc.add_paragraph(candidate.notes)
            doc.add_paragraph()

        # ── Details (status + source) ──
        self._section_heading(doc, "Details")
        doc.add_paragraph(
            f"Status: {candidate.get_status_display()}\n"
            f"Source: {candidate.get_source_display()}"
        )

        # ── Stream DOCX bytes ──
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{candidate.first_name}_{candidate.last_name}_CV.docx"
        content_type = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        )
        response = HttpResponse(buffer.getvalue(), content_type=content_type)
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    @staticmethod
    def _section_heading(doc, text):
        """Adds a bold blue section heading with a bottom border paragraph."""
        para = doc.add_paragraph()
        run  = para.add_run(text.upper())
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
